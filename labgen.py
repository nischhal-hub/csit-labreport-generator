import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types


class LabReportGenerator:
    def __init__(self, api_key: str, filename: str = "lab_report.docx"):
        self.client = genai.Client(api_key=api_key)
        self.filename = filename

    def get_genai_response(self, prompt: str) -> dict:
        """Fetch structured response from Gemini API."""
        json_schema = types.Schema(
            type="OBJECT",
            properties={
                "topic": types.Schema(type="STRING"),
                "objective": types.Schema(
                    type="ARRAY", items=types.Schema(type="STRING")
                ),
                "theory": types.Schema(type="STRING"),
                "implementation": types.Schema(type="STRING"),
                "conclusion": types.Schema(type="STRING"),
            },
        )
        specific_instruction = "The 'implementation' field should contain well-formatted, multi-line Java code with proper indentation."
        response = self.client.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=prompt + " " + specific_instruction,
            config=types.GenerateContentConfig(
                responseMimeType="application/json",
                responseSchema=json_schema,
                temperature=0.2,
            ),
        )
        return json.loads(response.text)

    def generate_document(self, data: dict):
        """Generate lab report as a .docx file."""
        doc = Document()

        # Set margins (Left=1.5, others=1)
        for section in doc.sections:
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # --- Formatting helpers ---
        def add_header(text: str):
            p = doc.add_paragraph(text.upper())
            run = p.runs[0]
            run.font.size = Pt(16)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def add_subheader(text: str):
            p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def add_paragraph(text: str):
            p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def add_code(text: str):
            p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # --- Start Writing Report ---
        add_header(f"TOPIC: {data['topic']}")

        add_subheader("1. OBJECTIVE")
        for obj in data.get("objective", []):
            obj_para = doc.add_paragraph(style="List Bullet")
            run = obj_para.add_run(obj)
            run.font.size = Pt(12)

        add_subheader("2. THEORY")
        add_paragraph(data.get("theory", ""))

        add_subheader("3. IMPLEMENTATION")
        add_code(data.get("implementation", ""))

        add_subheader("4. OUTPUT")
        add_paragraph("[Screenshot will be added here]")

        add_subheader("5. CONCLUSION")
        add_paragraph(data.get("conclusion", ""))

        # Save document
        doc.save(self.filename)
        print(f"✅ Lab report saved as {self.filename}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Lab Report using Gemini API")
    parser.add_argument("--api", help="Your Gemini API key")
    parser.add_argument("--prompt", help="Prompt for generating the lab report")
    parser.add_argument("--file", help="Output file name (default: lab_report.docx)")

    args = parser.parse_args()

    # Ask interactively if args are missing
    api_key = args.api or input("🔑 Enter your Gemini API key: ").strip() or "AIzaSyDuZGP7EDg1j522UKyErmNLzlSRBBgkQhk"
    prompt = args.prompt or input("📝 Enter your lab report prompt: ").strip()
    filename = args.file or input("📂 Enter output file name (default: lab_report.docx): ").strip() or "lab-report.docx"

    generator = LabReportGenerator(api_key=api_key, filename=filename)
    data = generator.get_genai_response(prompt)
    generator.generate_document(data)
